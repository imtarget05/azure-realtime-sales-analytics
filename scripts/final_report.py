"""
Final Report Generator — 100% based on verified source code.
Every claim is traceable to an actual file in the project.
"""
import os
import sys
import shutil
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "Report.docx")
BACKUP_PATH = os.path.join(os.path.dirname(__file__), "..", "Report_backup_final.docx")

# ── Helpers ──────────────────────────────────────────────────

def set_run(run, size=12, bold=False, italic=False, color=None, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text, style="Normal", size=12, bold=False, italic=False,
                  alignment=None, space_after=6, space_before=0):
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_bullet(doc, text, size=12, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run(r1, size=size, bold=True)
        r2 = p.add_run(text)
        set_run(r2, size=size)
    else:
        r = p.add_run(text)
        set_run(r, size=size)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, size=11, bold=True)
        # Light blue background
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "D6EAF8", qn("w:val"): "clear"
        })
        shading.append(bg)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run(run, size=11)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_code_block(doc, code, size=10):
    """Add a code-like paragraph with monospace font."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(code)
    set_run(run, size=size, name="Consolas", color=(40, 40, 40))
    return p


# ── MAIN ─────────────────────────────────────────────────────

def build_report():
    if os.path.exists(REPORT_PATH):
        shutil.copy2(REPORT_PATH, BACKUP_PATH)
        print(f"Backup: {BACKUP_PATH}")

    doc = Document()

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # ================================================================
    # TRANG BÌA
    # ================================================================
    add_paragraph(doc, "ĐẠI HỌC QUỐC GIA TP.HCM", size=14, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN", size=14, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "KHOA HỆ THỐNG THÔNG TIN", size=14, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    add_paragraph(doc, "BÀI TẬP 10", size=16, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_paragraph(doc, "HỆ THỐNG PHÂN TÍCH BÁN HÀNG THỜI GIAN THỰC", size=18, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_paragraph(doc, "TRÊN NỀN TẢNG MICROSOFT AZURE", size=18, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    add_paragraph(doc, "Môn: IS402 – Điện toán đám mây và ứng dụng", size=13,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc, "Giảng viên: ThS. Huỳnh Xuân Phụng", size=13,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_paragraph(doc, "Nhóm sinh viên thực hiện:", size=13, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    members = [
        ("Nguyễn Thái Bảo", "22520108"),
        ("Nguyễn Minh Phú", "22521100"),
        ("Bùi Quốc Huy", "22520530"),
    ]
    for name, mssv in members:
        add_paragraph(doc, f"{name} – {mssv}", size=13,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    add_paragraph(doc, f"TP. Hồ Chí Minh, tháng 4/2026", size=13,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, space_before=30)

    doc.add_page_break()

    # ================================================================
    # MỤC LỤC
    # ================================================================
    add_heading(doc, "MỤC LỤC", level=1)
    toc_items = [
        "Chương 1: Giới thiệu bài toán (1.5đ)",
        "  1.1 Loại bài toán",
        "  1.2 Loại dữ liệu",
        "  1.3 Kích thước dữ liệu và so sánh tốc độ",
        "  1.4 Phân loại dịch vụ cloud",
        "Chương 2: Cơ sở lý thuyết (1.5đ)",
        "  2.1 Định dạng lưu trữ",
        "  2.2 Thuật toán xử lý",
        "  2.3 Chi tiết dịch vụ cloud sử dụng",
        "Chương 3: Mô hình dữ liệu (2đ)",
        "  3.1 Tốc độ đọc ghi",
        "  3.2 Luồng xử lý ETL tự động",
        "  3.3 Sơ đồ quan hệ dữ liệu",
        "  3.4 Stored Procedures và Views",
        "Chương 4: Thực nghiệm và kết quả (3đ)",
        "  4.1 Kết quả benchmark",
        "  4.2 Machine Learning pipeline",
        "  4.3 Monitoring và MLOps",
        "Chương 5: Kết luận",
        "Tài liệu tham khảo",
    ]
    for item in toc_items:
        add_paragraph(doc, item, size=12, space_after=2)

    doc.add_page_break()

    # ================================================================
    # TÓM TẮT
    # ================================================================
    add_heading(doc, "TÓM TẮT", level=1)
    add_paragraph(doc,
        "Báo cáo trình bày hệ thống phân tích bán hàng thời gian thực trên Microsoft Azure, "
        "xử lý dữ liệu từ 3 cửa hàng (Hồ Chí Minh, Hà Nội, Đà Nẵng) với 35 sản phẩm thuộc 12 danh mục. "
        "Hệ thống sử dụng kiến trúc Lambda (Speed Layer + Batch Layer + Serving Layer) với 15 tài nguyên Azure "
        "được triển khai qua Terraform. Dữ liệu benchmark đạt 4.5 GB / 72.5 triệu dòng, "
        "mô hình GradientBoostingRegressor được chọn sau khi so sánh 9 thuật toán."
    )

    doc.add_page_break()

    # ================================================================
    # CHƯƠNG 1: GIỚI THIỆU BÀI TOÁN (1.5đ)
    # ================================================================
    add_heading(doc, "CHƯƠNG 1: GIỚI THIỆU BÀI TOÁN", level=1)

    # 1.1 Loại bài toán (0.25đ)
    add_heading(doc, "1.1 Loại bài toán (0.25đ)", level=2)
    add_paragraph(doc,
        "Hệ thống thuộc đồng thời 4 nhóm bài toán cloud computing:")
    add_table(doc,
        ["Nhóm", "Mô tả", "Minh chứng source code"],
        [
            ["Thu thập (Collection)",
             "Thu thập sự kiện bán hàng real-time qua Event Hub, làm giàu bằng OpenWeather API và Calendarific API",
             "data_generator/sales_generator.py\nconfig/settings.py: EVENT_HUB_NAME=\"sales-events\""],
            ["Lưu trữ (Storage)",
             "Azure SQL Database (SalesAnalyticsDB) lưu giao dịch, Blob Storage lưu reference data và model artifacts",
             "sql/create_tables.sql: 5 bảng + 7 indexes\nblob_storage/upload_reference_data.py: 3 containers"],
            ["Xử lý (Processing)",
             "Stream Analytics xử lý luồng real-time, Data Factory orchestrate ETL batch, ML training trên compute cluster",
             "stream_analytics/stream_query.sql\ndata_factory/pipeline_definition.json: 4 activities"],
            ["Trực quan (Visualization)",
             "Power BI dashboard 5 trang, Flask Web App với /predict và /dashboard",
             "powerbi/dashboard_layout.json: 5 pages\nwebapp/app.py: 10 routes"],
        ])

    # 1.2 Loại dữ liệu (0.25đ)
    add_heading(doc, "1.2 Loại dữ liệu (0.25đ)", level=2)
    add_paragraph(doc, "Hệ thống sử dụng nhiều loại dữ liệu:")
    add_table(doc,
        ["Loại", "Mô tả", "Nguồn trong source code"],
        [
            ["Database",
             "Azure SQL Database – SalesAnalyticsDB chứa 5 bảng chính + 1 bảng monitoring",
             "sql/create_tables.sql: SalesTransactions, HourlySalesSummary, SalesForecast, SalesAlerts, WeatherSalesCorrelation\nsql/create_monitoring_tables.sql: MonitoringEvents"],
            ["Dataset",
             "Dữ liệu CSV/Parquet sinh từ data generator (72.5M dòng, 4.5 GB)",
             "benchmarks/benchmark_data_size.py: generate_large_dataset()\nsample_events.jsonl"],
            ["Web API",
             "OpenWeather API (nhiệt độ, thời tiết), Calendarific API (ngày lễ), Stock API (giá cổ phiếu)",
             "data_generator/sales_generator.py: api.openweathermap.org, calendarific.com\ndata_generator/stock_generator.py"],
            ["ML Model",
             "GradientBoostingRegressor (scikit-learn) với 14 features, 2 targets",
             "ml/train_model.py: n_estimators=200, max_depth=6"],
        ])

    # 1.3 Kích thước dữ liệu (0.75đ)
    add_heading(doc, "1.3 Kích thước dữ liệu và so sánh tốc độ (0.75đ)", level=2)
    add_paragraph(doc,
        "Dữ liệu benchmark được sinh bởi benchmarks/benchmark_data_size.py, "
        "kết quả lưu tại benchmark_output/benchmark_report.json:", bold=True)

    add_paragraph(doc, "Thông số hệ thống benchmark:", bold=True, space_before=6)
    add_bullet(doc, "OS: Windows 10, CPU: 4 physical / 8 logical cores, RAM: 23.73 GB")
    add_bullet(doc, "Python: 3.11.0, thư viện: Polars (eager + lazy), PyArrow, pyodbc")

    add_paragraph(doc, "Kích thước dữ liệu:", bold=True, space_before=6)
    add_table(doc,
        ["Định dạng", "Dung lượng", "Số dòng", "Tỷ lệ nén"],
        [
            ["CSV", "4.5 GB", "72,500,000", "1.0x (gốc)"],
            ["Parquet", "1.3 GB", "72,500,000", "3.46x"],
        ])

    add_paragraph(doc, "So sánh tốc độ xử lý Local vs Cloud:", bold=True, space_before=8)
    add_paragraph(doc, "(Nguồn: benchmark_output/benchmark_report.json)", size=11, italic=True)
    add_table(doc,
        ["Phương pháp", "Thời gian", "Throughput", "Ghi chú"],
        [
            ["Local CSV (Polars)", "123.03s", "37.49 MB/s", "I/O chiếm 92.7%"],
            ["Local Parquet (eager)", "3.86s", "461.48 MB/s", "Nhanh hơn CSV 31.87×"],
            ["Local Parquet (lazy)", "2.33s", "—", "Tối ưu bộ nhớ"],
            ["Cloud SQL (500K rows) – COUNT(*)", "1.37s", "—", ""],
            ["Cloud SQL (500K rows) – SUM", "1.98s", "—", ""],
            ["Cloud SQL (500K rows) – GROUP BY store", "3.34s", "—", ""],
            ["Cloud SQL (500K rows) – GROUP BY category", "3.84s", "—", ""],
            ["Cloud SQL (500K rows) – TOP 10", "3.34s", "—", ""],
            ["Cloud SQL tổng 5 queries", "13.87s", "—", ""],
        ])

    add_paragraph(doc,
        "So sánh công bằng (cùng 500K rows): Local Parquet = 0.094s vs Cloud SQL = 13.87s. "
        "Cloud chậm hơn 147× cho batch analytics do network latency. "
        "Tuy nhiên, Cloud vượt trội cho real-time concurrent queries "
        "(trung bình 27.67ms/query đơn — benchmark_output/benchmark_latency.json) "
        "và auto-scaling cho hàng trăm kết nối đồng thời.",
        space_before=6)

    add_paragraph(doc, "Latency đo từ Việt Nam đến các Azure regions:", bold=True, space_before=6)
    add_paragraph(doc, "(Nguồn: benchmark_output/benchmark_latency.json)", size=11, italic=True)
    add_table(doc,
        ["Region", "Avg Latency", "Ghi chú"],
        [
            ["Southeast Asia", "69.28 ms", "Tốt nhất – được khuyến nghị"],
            ["Japan East", "130.49 ms", ""],
            ["Australia East", "149.56 ms", ""],
            ["West Europe", "216.13 ms", ""],
            ["East US", "290.73 ms", "Hiện tại đang dùng"],
        ])

    # 1.4 Phân loại dịch vụ cloud (0.25đ)
    add_heading(doc, "1.4 Phân loại dịch vụ cloud (0.25đ)", level=2)
    add_paragraph(doc,
        "Phân loại theo mô hình dịch vụ (Nguồn: terraform/main.tf, 15 resources):")
    add_table(doc,
        ["Phân loại", "Dịch vụ", "Resource trong Terraform"],
        [
            ["IaaS", "Azure ML Compute Cluster (Standard_DS3_v2, max 4 nodes)",
             "azurerm_machine_learning_compute_cluster.training"],
            ["PaaS", "Azure SQL Database (SKU S0)",
             "azurerm_mssql_database.main"],
            ["PaaS", "Azure Event Hub (Standard, 4 partitions)",
             "azurerm_eventhub.sales"],
            ["PaaS", "Azure Stream Analytics (3 SU)",
             "azurerm_stream_analytics_job.main"],
            ["PaaS", "Azure Blob Storage (Standard LRS)",
             "azurerm_storage_account.main"],
            ["PaaS", "Azure Key Vault (Standard, soft-delete 7 ngày)",
             "azurerm_key_vault.main"],
            ["PaaS", "Azure ML Workspace",
             "azurerm_machine_learning_workspace.main"],
            ["PaaS", "Azure App Service (Web App)",
             "Dockerfile + webapp/app.py"],
            ["PaaS", "Azure Data Factory",
             "data_factory/pipeline_definition.json"],
            ["PaaS", "Application Insights",
             "azurerm_application_insights.main"],
            ["FaaS", "Azure Functions (ValidateSalesEvent, DriftMonitor)",
             "azurerm_linux_function_app.validation (Y1 Consumption)"],
            ["SaaS", "Power BI (Dashboard + Streaming Dataset)",
             "powerbi/dashboard_layout.json, powerbi/dax_measures.dax"],
        ])

    doc.add_page_break()

    # ================================================================
    # CHƯƠNG 2: CƠ SỞ LÝ THUYẾT (1.5đ)
    # ================================================================
    add_heading(doc, "CHƯƠNG 2: CƠ SỞ LÝ THUYẾT", level=1)

    # 2.1 Định dạng lưu trữ (0.5đ)
    add_heading(doc, "2.1 Định dạng lưu trữ (0.5đ)", level=2)

    add_heading(doc, "2.1.1 Azure SQL Database (Relational)", level=3)
    add_paragraph(doc,
        "Database: SalesAnalyticsDB (config/settings.py: SQL_DATABASE=\"SalesAnalyticsDB\")")
    add_paragraph(doc, "Schema gồm 6 bảng (sql/create_tables.sql + sql/create_monitoring_tables.sql):", bold=True)
    add_table(doc,
        ["Bảng", "Mục đích", "Cột chính"],
        [
            ["SalesTransactions", "Giao dịch raw từ Stream Analytics",
             "event_time, store_id, product_id, units_sold, unit_price, revenue, temperature, weather, holiday, category, ingest_lag_seconds"],
            ["HourlySalesSummary", "Aggregation 5-phút (TumblingWindow)",
             "window_start, window_end, store_id, product_id, units_sold, revenue, avg_price, tx_count, rolling_15m_units/revenue"],
            ["SalesForecast", "Dự đoán ML",
             "forecast_date, forecast_hour, store_id, predicted_quantity, predicted_revenue, confidence_lower/upper, model_version"],
            ["SalesAlerts", "Cảnh báo anomaly từ AnomalyDetection_SpikeAndDip",
             "alert_time, store_id, type, value"],
            ["WeatherSalesCorrelation", "Tương quan thời tiết – doanh thu",
             "window_end, store_id, weather, avg_temperature, avg_stock_price, total_revenue, correlation_signal"],
            ["MonitoringEvents", "Log monitoring drift detection",
             "event_type, mae_value, threshold, model_version, retrain_triggered, details (JSON)"],
        ])

    add_paragraph(doc, "Indexes (7 indexes trong create_tables.sql):", bold=True, space_before=6)
    add_bullet(doc, "IX_SalesTransactions_EventTime, IX_SalesTransactions_StoreProduct")
    add_bullet(doc, "IX_HourlySalesSummary_Window, IX_SalesForecast_Date, IX_SalesForecast_DateTime")
    add_bullet(doc, "IX_SalesAlerts_TimeStore, IX_WeatherSalesCorrelation_WindowStore")

    add_paragraph(doc, "Views (9 views tổng cộng):", bold=True, space_before=6)
    add_bullet(doc, "create_tables.sql: vw_RealtimeDashboard (TOP 1000 recent), vw_ForecastVsActual (forecast vs actual)")
    add_bullet(doc, "create_powerbi_views.sql: vw_ETLPipelineHealth, vw_ProductPerformance, vw_StoreComparison, vw_WeatherImpact, vw_AlertSummary, vw_ForecastAccuracy, vw_HourlyTrend")
    add_bullet(doc, "create_monitoring_tables.sql: vw_MonitoringSummary")

    add_heading(doc, "2.1.2 Azure Blob Storage (Object)", level=3)
    add_paragraph(doc,
        "4 containers (terraform/main.tf: azurerm_storage_container.containers):")
    add_table(doc,
        ["Container", "Mục đích", "Nội dung"],
        [
            ["reference-data", "Reference data cho Stream Analytics",
             "products.json, regions.json, customer_segments.json, payment_methods.json"],
            ["sales-archive", "Archive từ Event Hub Capture",
             "Historical data backup, retrain artifacts"],
            ["data-factory-staging", "Staging cho ADF pipeline",
             "JSON data cho CopyBlobToSQL activity"],
            ["ml-artifacts", "Model artifacts",
             "revenue_model.pkl, quantity_model.pkl, label_encoders.pkl, model_metadata.json"],
        ])

    add_heading(doc, "2.1.3 Event Hub (Streaming)", level=3)
    add_paragraph(doc, "3 Event Hubs (terraform/main.tf):")
    add_table(doc,
        ["Event Hub", "Partitions", "Retention", "Mục đích"],
        [
            ["sales-events", "4", "7 ngày", "Sự kiện bán hàng chính"],
            ["weather-events", "2", "7 ngày", "Dữ liệu thời tiết"],
            ["stock-events", "2", "7 ngày", "Dữ liệu giá cổ phiếu"],
        ])

    add_heading(doc, "2.1.4 Model Artifacts", level=3)
    add_paragraph(doc, "Định dạng: pickle (.pkl) – scikit-learn serialization")
    add_bullet(doc, "revenue_model.pkl — GradientBoostingRegressor cho dự đoán doanh thu")
    add_bullet(doc, "quantity_model.pkl — GradientBoostingRegressor cho dự đoán số lượng")
    add_bullet(doc, "label_encoders.pkl — LabelEncoder cho store_id, product_id, category")
    add_bullet(doc, "model_metadata.json — metrics (MAE, RMSE, R²), feature_columns, trained_at, model_version")

    # 2.2 Thuật toán xử lý (0.5đ)
    add_heading(doc, "2.2 Thuật toán xử lý (0.5đ)", level=2)

    add_heading(doc, "2.2.1 Stream Processing (stream_analytics/stream_query.sql)", level=3)
    add_paragraph(doc, "Pipeline xử lý gồm 5 CTE (Common Table Expression) nối tiếp:")

    add_paragraph(doc, "Giai đoạn 1 — Cleaned: Lọc NULL, ép kiểu", bold=True)
    add_code_block(doc,
        "FROM SalesInput TIMESTAMP BY [timestamp]\n"
        "WHERE TRY_CAST([timestamp] AS datetime) IS NOT NULL\n"
        "  AND store_id IS NOT NULL AND product_id IS NOT NULL")

    add_paragraph(doc, "Giai đoạn 2 — Enriched: Tính revenue, phân loại category", bold=True)
    add_code_block(doc,
        "revenue = quantity × price\n"
        "Category mapping: 35 product_ids → 12 categories\n"
        "  (Beverage, Dairy, Bakery, Electronics, Clothing, Home,\n"
        "   Accessories, Snacks, Health & Beauty, Sports, Stationery, Toys)")

    add_paragraph(doc, "Giai đoạn 3 — Agg5m: TumblingWindow(minute, 5)", bold=True)
    add_code_block(doc,
        "GROUP BY store_id, product_id, category, TumblingWindow(minute, 5)\n"
        "→ SUM(units_sold), SUM(revenue), AVG(unit_price), COUNT(*)")

    add_paragraph(doc, "Giai đoạn 4 — AnomalyDetection: Phát hiện bất thường", bold=True)
    add_code_block(doc,
        "AnomalyDetection_SpikeAndDip(\n"
        "  CAST(revenue AS bigint), 95, 120, 'spikesanddips'\n"
        ") OVER (PARTITION BY store_id LIMIT DURATION(minute, 30))\n"
        "→ Confidence: 95%, Lookback: 120 samples, Window: 30 phút")

    add_paragraph(doc, "Giai đoạn 5 — Output: 4 luồng song song", bold=True)
    add_table(doc,
        ["Output", "Đích", "Window"],
        [
            ["SalesTransactionsOutput", "dbo.SalesTransactions", "Không (raw events)"],
            ["HourlySalesSummaryOutput", "dbo.HourlySalesSummary", "TumblingWindow(minute, 5)"],
            ["SalesAlertsOutput", "dbo.SalesAlerts", "Chỉ anomaly (IsAnomaly=1)"],
            ["PowerBIRealtimeOutput", "Power BI Streaming Dataset", "TumblingWindow(minute, 1)"],
        ])

    add_paragraph(doc,
        "Lưu ý: ARM template (infrastructure/arm_streaming_job.json) còn có HoppingWindow(minute, 15, 5) "
        "cho rolling 15-phút aggregation và LAG function cho revenue_delta_5m.",
        italic=True, space_before=6)

    add_heading(doc, "2.2.2 Machine Learning (ml/train_model.py)", level=3)
    add_paragraph(doc, "Thuật toán chính: GradientBoostingRegressor (scikit-learn)", bold=True)
    add_table(doc,
        ["Tham số", "Giá trị", "Nguồn"],
        [
            ["n_estimators", "200", "ml/train_model.py dòng 195"],
            ["max_depth", "6", "ml/train_model.py dòng 195"],
            ["learning_rate", "0.1", "ml/train_model.py dòng 195"],
            ["subsample", "0.8", "ml/train_model.py dòng 195"],
            ["min_samples_split", "10", "ml/train_model.py dòng 195"],
            ["min_samples_leaf", "5", "ml/train_model.py dòng 195"],
            ["random_state", "42", "ml/train_model.py dòng 195"],
            ["test_size", "0.2", "train_test_split"],
        ])

    add_paragraph(doc, "Feature Engineering (14 features):", bold=True, space_before=6)
    add_table(doc,
        ["Nhóm", "Features", "Kỹ thuật"],
        [
            ["Temporal", "hour, day_of_month, month, is_weekend", "Trích từ timestamp"],
            ["Cyclic", "hour_sin, hour_cos, month_sin, month_cos",
             "sin(2π·hour/24), cos(2π·hour/24), sin(2π·month/12), cos(2π·month/12)"],
            ["Location", "store_id_enc", "LabelEncoder"],
            ["Product", "product_id_enc, category_enc", "LabelEncoder"],
            ["Environment", "temperature, is_rainy, holiday", "Số thực / binary"],
        ])

    add_paragraph(doc, "2 Target riêng biệt: quantity (model 1) và revenue (model 2)", bold=True)

    add_heading(doc, "2.2.3 So sánh 9 mô hình (ml/compare_models.py)", level=3)
    add_paragraph(doc,
        "So sánh 9 thuật toán trên 20,000 mẫu synthetic data, 3-fold cross-validation:")
    add_table(doc,
        ["#", "Model", "Tham số chính", "Scaling"],
        [
            ["1", "Linear Regression", "default", "StandardScaler"],
            ["2", "Ridge Regression", "alpha=1.0", "StandardScaler"],
            ["3", "Lasso Regression", "alpha=0.1", "StandardScaler"],
            ["4", "Decision Tree", "max_depth=10", "Không"],
            ["5", "Random Forest", "n_estimators=100, max_depth=10", "Không"],
            ["6", "Gradient Boosting", "n_estimators=200, max_depth=5", "Không"],
            ["7", "AdaBoost", "n_estimators=100", "Không"],
            ["8", "KNN (k=5)", "n_neighbors=5", "StandardScaler"],
            ["9", "SVR (RBF)", "kernel=rbf, C=100", "StandardScaler"],
        ])
    add_paragraph(doc,
        "Kết quả chi tiết xuất tại: benchmark_output/ml_comparison/. "
        "GradientBoosting được chọn vì: cân bằng giữa accuracy và training time, "
        "xử lý tốt non-linear relationships trên tabular data, "
        "hỗ trợ feature importance để giải thích mô hình.",
        space_before=4)

    add_heading(doc, "2.2.4 Drift Detection (ml/drift_monitor.py)", level=3)
    add_paragraph(doc,
        "Thuật toán: MAE-based drift detection (Mean Absolute Error)")
    add_bullet(doc, "Nguồn dữ liệu: SQL view dbo.vw_ForecastVsActual — so sánh forecast vs actual revenue")
    add_bullet(doc, "Threshold mặc định: MAE > 25.0 (config/settings.py: DRIFT_MAE_ABS_THRESHOLD=25.0)")
    add_bullet(doc, "Window: 24 giờ gần nhất, tối thiểu 24 mẫu")
    add_bullet(doc, "Cooldown: 120 phút giữa 2 lần trigger")
    add_bullet(doc, "Trigger modes: local (retrain_and_compare.py), azureml (mlops/trigger_training_pipeline.py), both, none")
    add_bullet(doc, "Concurrency: File-based lock với PID tracking, timeout 600s")

    # 2.3 Chi tiết dịch vụ cloud (0.5đ)
    add_heading(doc, "2.3 Chi tiết dịch vụ cloud sử dụng (0.5đ)", level=2)

    services = [
        ("Azure Event Hub", "PaaS",
         "Message broker nhận sự kiện bán hàng real-time từ data generator",
         "Namespace Standard, capacity 1. Event Hub 'sales-events': 4 partitions, 7 ngày retention. "
         "Consumer group 'stream-analytics-cg' cho ASA. Transport: AmqpOverWebsocket.",
         "terraform/main.tf: azurerm_eventhub_namespace, config/settings.py: EVENT_HUB_NAME"),

        ("Azure Stream Analytics", "PaaS",
         "Stream processing engine — cleaning, enrichment, aggregation, anomaly detection",
         "3 Streaming Units, compatibility level 1.2. Late arrival: 60s, out-of-order: 50s. "
         "Input: SalesInput (Event Hub). Outputs: 4 (SQL×3 + Power BI×1). "
         "Sử dụng TumblingWindow(5min, 1min) và AnomalyDetection_SpikeAndDip.",
         "terraform/main.tf: azurerm_stream_analytics_job, stream_analytics/stream_query.sql"),

        ("Azure SQL Database", "PaaS",
         "Relational database lưu trữ giao dịch, aggregation, forecast, alerts, monitoring",
         "Server version 12.0, SKU S0, TLS 1.2. Database: SalesAnalyticsDB. "
         "6 bảng, 7 indexes, 10 views, 3 stored procedures.",
         "terraform/main.tf: azurerm_mssql_database, sql/create_tables.sql"),

        ("Azure Data Factory", "PaaS",
         "ETL orchestration — batch processing pipeline",
         "Pipeline 'SalesAnalyticsPipeline' với 4 activities tuần tự: "
         "CopyBlobToSQL (Copy) → PrepareTrainingData (SP) → RunMLPipeline (AzureMLExecutePipeline) → "
         "UpdateForecasts (SP). Trigger: DailyMLTrigger chạy 02:00 UTC hàng ngày.",
         "data_factory/pipeline_definition.json"),

        ("Azure Machine Learning", "PaaS + IaaS (Compute)",
         "MLOps platform — training, registry, endpoint deployment",
         "Workspace: aml-sales-analytics. Compute cluster: 'training-cluster' (Dedicated, max 4 nodes, "
         "auto-scale down sau 15 phút idle). Model: 'sales-forecast-model'. "
         "Endpoint: 'sales-forecast-endpoint' (Managed Online, blue/green deployment).",
         "terraform/main.tf: azurerm_machine_learning_workspace, mlops/trigger_training_pipeline.py, mlops/deploy_to_endpoint.py"),

        ("Azure Blob Storage", "PaaS",
         "Object storage cho reference data, model artifacts, archive",
         "Account Standard LRS, TLS 1.2. 4 containers: reference-data, sales-archive, "
         "data-factory-staging, ml-artifacts.",
         "terraform/main.tf: azurerm_storage_account, blob_storage/upload_reference_data.py"),

        ("Azure Functions", "FaaS",
         "Serverless event validation và drift monitoring",
         "Function 1: ValidateSalesEvent (Event Hub trigger) — validate fields, types, ranges, dedup. "
         "Function 2: DriftMonitor (Timer trigger mỗi 1 giờ) — gọi ml/drift_monitor.py. "
         "Runtime: Python 3.10, SKU Y1 Consumption.",
         "azure_functions/ValidateSalesEvent/__init__.py, azure_functions/DriftMonitor/__init__.py"),

        ("Azure Key Vault", "PaaS",
         "Secret management — lưu trữ connection strings, API keys",
         "Standard SKU, soft-delete 7 ngày. 9 secrets: event-hub-connection-string, sql-connection-string, "
         "sql-admin-password, ml-endpoint-url, ml-api-key, blob-connection-string, powerbi-push-url, "
         "appinsights-connection-string, openweather-api-key.",
         "security/key_vault.py: SecretManager class, terraform/main.tf: azurerm_key_vault"),

        ("Application Insights", "PaaS",
         "Monitoring, telemetry, distributed tracing",
         "SDK: opencensus. Tracks: custom events, metrics, dependencies, exceptions. "
         "PipelineHealthMonitor theo dõi 7 components. Decorator @monitor_performance tự động đo latency.",
         "monitoring/telemetry.py, terraform/main.tf: azurerm_application_insights"),

        ("Power BI", "SaaS",
         "Dashboard trực quan hóa KPI bán hàng",
         "5 trang: Overview, Products, Customers, Anomaly, AccessRights. "
         "Canvas 1280×720 + mobile 360×640. Theme dark. "
         "RLS: 3 roles (RegionManager, Admin, AccessRightAdmin) qua SecurityMapping table. "
         "30+ DAX measures. Streaming dataset từ ASA PowerBIRealtimeOutput.",
         "powerbi/dashboard_layout.json, powerbi/dax_measures.dax, powerbi/rls_config.dax"),

        ("Azure App Service", "PaaS",
         "Hosting Flask web application",
         "Framework: Flask. 10 routes: /, /predict, /api/predict, /api/ingest, "
         "/api/health, /dashboard, /model-report, /api/benchmark/<name>. "
         "ML scoring cascade: Azure ML Endpoint → Local pkl → Mock fallback.",
         "webapp/app.py, Dockerfile"),

        ("Azure Databricks", "PaaS",
         "Advanced analytics platform",
         "Host: adb-7405607469187602.2.azuredatabricks.net. "
         "SQL scripts cho access rights và security mapping.",
         "databricks/sql/create_access_rights.sql, databricks/sql/create_security_mapping.sql"),
    ]

    for svc_name, svc_type, role, details, source in services:
        add_paragraph(doc, f"{svc_name} ({svc_type})", size=13, bold=True, space_before=8)
        add_bullet(doc, f"Vai trò: {role}")
        add_bullet(doc, f"Cấu hình: {details}")
        add_bullet(doc, f"Source: {source}")

    doc.add_page_break()

    # ================================================================
    # CHƯƠNG 3: MÔ HÌNH DỮ LIỆU (2đ)
    # ================================================================
    add_heading(doc, "CHƯƠNG 3: MÔ HÌNH DỮ LIỆU", level=1)

    # 3.1 Tốc độ đọc ghi (0.5đ)
    add_heading(doc, "3.1 Tốc độ đọc ghi (0.5đ)", level=2)
    add_paragraph(doc,
        "Benchmark đo bằng benchmarks/benchmark_read_write.py, kết quả tại benchmark_output/benchmark_read_write.json:",
        bold=True)

    add_paragraph(doc, "Write Performance:", bold=True, space_before=6)
    add_table(doc,
        ["Thao tác", "Throughput", "Latency/row"],
        [
            ["Single INSERT", "~80 rows/s", "~12.45ms/row"],
            ["Batch INSERT (1000/batch)", "~1,250 rows/s", "~0.87ms/row"],
        ])
    add_paragraph(doc, "→ Batch INSERT nhanh hơn Single INSERT 15.6 lần", bold=True)

    add_paragraph(doc, "Read Performance:", bold=True, space_before=6)
    add_table(doc,
        ["Query", "Latency"],
        [
            ["COUNT(*)", "8.9 ms"],
            ["SUM(revenue)", "15.6 ms"],
            ["GROUP BY store_id", "34.2 ms"],
            ["GROUP BY category", "48.7 ms"],
            ["JOIN + SUM", "62.3 ms"],
            ["Window function ranking", "89.1 ms"],
            ["Multi-condition filter", "23.4 ms"],
        ])

    add_paragraph(doc, "Network Latency đến SQL Server:", bold=True, space_before=6)
    add_paragraph(doc, "(Nguồn: benchmark_output/benchmark_latency.json)", size=11, italic=True)
    add_table(doc,
        ["Metric", "Giá trị"],
        [
            ["TCP latency (avg)", "40.54 ms"],
            ["SQL connection time", "433.76 ms"],
            ["Simple query latency (avg)", "27.67 ms"],
            ["Complex query latency (avg)", "26.82 ms"],
        ])

    # 3.2 Luồng xử lý ETL tự động (0.5đ)
    add_heading(doc, "3.2 Luồng xử lý ETL tự động (0.5đ)", level=2)

    add_paragraph(doc, "Luồng dữ liệu end-to-end (trích từ source code):", bold=True)
    add_paragraph(doc, "SPEED LAYER (Real-time):", bold=True, space_before=6)
    add_code_block(doc,
        "Data Generator (sales_generator.py)\n"
        "  → Azure Event Hub (sales-events, 4 partitions)\n"
        "    → Azure Functions/ValidateSalesEvent (validate, dedup)\n"
        "      → Azure Stream Analytics (stream_query.sql)\n"
        "        ├→ dbo.SalesTransactions (raw events)\n"
        "        ├→ dbo.HourlySalesSummary (TumblingWindow 5min)\n"
        "        ├→ dbo.SalesAlerts (AnomalyDetection)\n"
        "        └→ Power BI Streaming Dataset (TumblingWindow 1min)")

    add_paragraph(doc, "BATCH LAYER (Offline — hàng ngày 02:00 UTC):", bold=True, space_before=6)
    add_code_block(doc,
        "ADF Pipeline: SalesAnalyticsPipeline\n"
        "  1. CopyBlobToSQL (Blob staging → dbo.SalesTransactions)\n"
        "  2. PrepareTrainingData (sp_PrepareTrainingData → dbo.MLTrainingData)\n"
        "  3. RunMLPipeline (AzureMLExecutePipeline → train_and_register.py)\n"
        "  4. UpdateForecasts (sp_UpdateForecasts → clean old forecasts)")

    add_paragraph(doc, "SERVING LAYER:", bold=True, space_before=6)
    add_code_block(doc,
        "Azure ML Endpoint (score.py)\n"
        "  → Flask Web App (webapp/app.py)\n"
        "    ├→ /predict (dự đoán doanh thu)\n"
        "    ├→ /dashboard (monitoring real-time)\n"
        "    └→ /api/predict (REST API)\n"
        "Power BI Dashboard (5 trang, 30+ DAX measures, RLS 3 roles)")

    add_paragraph(doc, "MONITORING LAYER:", bold=True, space_before=6)
    add_code_block(doc,
        "DriftMonitor Azure Function (timer mỗi 1 giờ)\n"
        "  → ml/drift_monitor.py (MAE > 25.0?)\n"
        "    ├→ trigger_mode=local: ml/retrain_and_compare.py --promote\n"
        "    ├→ trigger_mode=azureml: mlops/trigger_training_pipeline.py\n"
        "    ├→ trigger_github_actions: .github/workflows/ci-cd-mlops.yml\n"
        "    └→ Alerts: Slack webhook + Email SMTP + Teams Adaptive Card\n"
        "GitHub Actions:\n"
        "  drift-detection.yml (cron mỗi 6 giờ)\n"
        "    → ci-cd-mlops.yml (Terraform + retrain + deploy)")

    # 3.3 Sơ đồ quan hệ (0.5đ)
    add_heading(doc, "3.3 Sơ đồ quan hệ dữ liệu (0.5đ)", level=2)
    add_paragraph(doc, "Quan hệ giữa các bảng SQL (sql/create_tables.sql):", bold=True)
    add_code_block(doc,
        "SalesTransactions\n"
        "  ├─ store_id ──┐\n"
        "  ├─ product_id ─┤── JOIN → HourlySalesSummary (cùng store_id, product_id)\n"
        "  └─ category ───┘\n"
        "\n"
        "SalesForecast\n"
        "  ├─ store_id, category ── JOIN → SalesTransactions\n"
        "  └─ → vw_ForecastVsActual (forecast vs actual aggregation)\n"
        "\n"
        "SalesAlerts\n"
        "  └─ store_id ── JOIN → SalesTransactions\n"
        "\n"
        "WeatherSalesCorrelation\n"
        "  └─ store_id, weather ── JOIN → SalesTransactions\n"
        "\n"
        "MonitoringEvents\n"
        "  └─ Standalone table (MLOps audit log)")

    add_paragraph(doc, "Power BI Data Model (powerbi/rls_config.dax):", bold=True, space_before=6)
    add_code_block(doc,
        "SecurityMapping[allowed_region] ──M:1──→ hourly_summary[region]\n"
        "  (Cross-filter: Both Directions)\n"
        "\n"
        "RLS Roles:\n"
        "  RegionManager: WHERE [user_email] = USERPRINCIPALNAME()\n"
        "  Admin: Không lọc (thấy toàn bộ)\n"
        "  AccessRightAdmin: Không lọc trên access_rights")

    # 3.4 Stored Procedures và Views (0.5đ)
    add_heading(doc, "3.4 Stored Procedures và Views (0.5đ)", level=2)

    add_paragraph(doc, "Stored Procedures:", bold=True)
    add_table(doc,
        ["Procedure", "File", "Chức năng"],
        [
            ["sp_PrepareTrainingData", "sql/stored_procedures.sql",
             "Tạo bảng MLTrainingData từ 90 ngày gần nhất SalesTransactions, GROUP BY date/hour/store/product/category"],
            ["sp_UpdateForecasts", "sql/stored_procedures.sql",
             "Xoá forecast cũ hơn 7 ngày từ SalesForecast"],
            ["sp_CleanupMonitoringEvents", "sql/create_monitoring_tables.sql",
             "Xoá monitoring events cũ hơn 90 ngày (tham số @RetentionDays)"],
        ])

    add_paragraph(doc, "Views (10 tổng cộng):", bold=True, space_before=6)
    add_table(doc,
        ["View", "File", "Mô tả"],
        [
            ["vw_RealtimeDashboard", "create_tables.sql", "TOP 1000 giao dịch gần nhất"],
            ["vw_ForecastVsActual", "create_tables.sql", "So sánh forecast vs actual (JOIN)"],
            ["vw_MonitoringSummary", "create_monitoring_tables.sql", "7 ngày gần nhất: event count, avg/max MAE"],
            ["vw_ETLPipelineHealth", "create_powerbi_views.sql", "Health pipeline theo 5-phút buckets"],
            ["vw_ProductPerformance", "create_powerbi_views.sql", "Performance per product 24h"],
            ["vw_StoreComparison", "create_powerbi_views.sql", "So sánh cửa hàng 7 ngày"],
            ["vw_WeatherImpact", "create_powerbi_views.sql", "Tương quan thời tiết-doanh thu"],
            ["vw_AlertSummary", "create_powerbi_views.sql", "Tổng hợp alert theo severity"],
            ["vw_ForecastAccuracy", "create_powerbi_views.sql", "Forecast accuracy + confidence"],
            ["vw_HourlyTrend", "create_powerbi_views.sql", "Trend 3 ngày từ HourlySalesSummary"],
        ])

    doc.add_page_break()

    # ================================================================
    # CHƯƠNG 4: THỰC NGHIỆM VÀ KẾT QUẢ (3đ)
    # ================================================================
    add_heading(doc, "CHƯƠNG 4: THỰC NGHIỆM VÀ KẾT QUẢ", level=1)

    # 4.1 Kết quả benchmark (1đ)
    add_heading(doc, "4.1 Kết quả benchmark (1đ)", level=2)

    add_heading(doc, "4.1.1 Benchmark dữ liệu lớn (benchmark_data_size.py)", level=3)
    add_paragraph(doc,
        "Sinh 72,500,000 dòng × 10 cột, lưu CSV 4.5 GB + chuyển đổi Parquet 1.3 GB.")

    add_paragraph(doc, "So sánh đọc dữ liệu:", bold=True, space_before=6)
    add_table(doc,
        ["Phương pháp", "Thời gian đọc", "Thời gian tính", "Tổng", "Throughput"],
        [
            ["Local CSV (Polars)", "114.03s", "9.01s", "123.03s", "37.49 MB/s"],
            ["Local Parquet (eager)", "2.88s", "0.974s", "3.86s", "461.48 MB/s"],
            ["Local Parquet (lazy)", "—", "—", "2.33s", "—"],
        ])
    add_paragraph(doc, "→ Parquet nhanh hơn CSV 31.87× nhờ columnar format + compression", bold=True)

    add_paragraph(doc, "Cloud SQL Queries (trên 500,000 rows):", bold=True, space_before=6)
    add_table(doc,
        ["Query", "Thời gian"],
        [
            ["SELECT COUNT(*)", "1.37s"],
            ["SELECT SUM(revenue)", "1.98s"],
            ["GROUP BY store_id", "3.34s"],
            ["GROUP BY category", "3.84s"],
            ["TOP 10 products", "3.34s"],
            ["Tổng 5 queries", "13.87s"],
        ])

    add_heading(doc, "4.1.2 Benchmark read/write (benchmark_read_write.py)", level=3)
    add_paragraph(doc, "Kết quả từ benchmark_output/benchmark_read_write.json:")
    add_bullet(doc, "Single INSERT: ~80 rows/s (12.45ms/row)")
    add_bullet(doc, "Batch INSERT (1000/batch): ~1,250 rows/s → nhanh hơn 15.6×")
    add_bullet(doc, "Query đơn: 8.9ms (COUNT) đến 89.1ms (window function)")

    add_heading(doc, "4.1.3 Benchmark latency (benchmark_latency.py)", level=3)
    add_paragraph(doc, "Kết quả từ benchmark_output/benchmark_latency.json:")
    add_bullet(doc, "DNS resolution: SQL Server 75.34ms, Azure Portal 33.86ms")
    add_bullet(doc, "TCP to SQL: avg 40.54ms (min 24.89, max 53.23)")
    add_bullet(doc, "SQL connection: 433.76ms")
    add_bullet(doc, "Simple query: avg 27.67ms (P50: 26.26ms)")
    add_bullet(doc, "Complex query: avg 26.82ms")
    add_bullet(doc, "Region tốt nhất: Southeast Asia (69.28ms) — khuyến nghị migrate")

    # 4.2 ML Pipeline (1đ)
    add_heading(doc, "4.2 Machine Learning Pipeline (1đ)", level=2)

    add_heading(doc, "4.2.1 Training (ml/train_model.py)", level=3)
    add_paragraph(doc,
        "Dữ liệu: 50,000 mẫu synthetic (fallback) hoặc dữ liệu SQL thực (nếu ≥1000 dòng). "
        "Mô hình: GradientBoostingRegressor(n_estimators=200, max_depth=6, learning_rate=0.1, "
        "subsample=0.8, min_samples_split=10, min_samples_leaf=5). "
        "Train 2 model riêng: revenue_model và quantity_model. "
        "Đánh giá: MAE, RMSE, R², 5-fold cross-validation R².")

    add_paragraph(doc, "Biểu đồ tự động sinh:", bold=True, space_before=4)
    add_bullet(doc, "Feature Importance (top features)")
    add_bullet(doc, "Actual vs Predicted scatter plot")
    add_bullet(doc, "Residual distribution")
    add_bullet(doc, "Learning curve")
    add_bullet(doc, "Error by hour of day")

    add_heading(doc, "4.2.2 So sánh 9 mô hình (ml/compare_models.py)", level=3)
    add_paragraph(doc,
        "So sánh trên 20,000 mẫu, 3-fold CV. 9 thuật toán: "
        "Linear Regression, Ridge(α=1.0), Lasso(α=0.1), Decision Tree(depth=10), "
        "Random Forest(100 trees, depth=10), Gradient Boosting(200 trees, depth=5), "
        "AdaBoost(100 trees), KNN(k=5), SVR(RBF, C=100).")
    add_paragraph(doc,
        "SVR giới hạn 10,000 mẫu training. SVR/KNN/Linear/Ridge/Lasso dùng StandardScaler. "
        "Output: matplotlib 4-panel comparison + plotly interactive HTML + radar chart + JSON ranking. "
        "Kết quả xuất tại benchmark_output/ml_comparison/.",
        space_before=2)

    add_heading(doc, "4.2.3 Retrain & Compare (ml/retrain_and_compare.py)", level=3)
    add_paragraph(doc,
        "So sánh old model (Ridge, α=1.0) vs new model (GradientBoosting, configurable: "
        "default n_estimators=300, max_depth=6, lr=0.1, subsample=0.8). "
        "Gate check: PROMOTE chỉ khi new R² ≥ old R² cho CẢ revenue VÀ quantity. "
        "History lưu tại model_output/retrain_history/ (giữ 50 runs gần nhất).")

    add_heading(doc, "4.2.4 Scoring & Deployment (ml/score.py, mlops/deploy_to_endpoint.py)", level=3)
    add_paragraph(doc,
        "score.py: Azure ML Endpoint scoring script. Input: 10 features → encode categorical → "
        "cyclic transform → predict → output: predicted_revenue, predicted_quantity, "
        "confidence interval (±1.96×RMSE). "
        "deploy_to_endpoint.py: Managed Online Endpoint, blue/green deployment với traffic splitting, "
        "liveness probe (30s initial, 10s period), readiness probe.")

    # 4.3 Monitoring & MLOps (1đ)
    add_heading(doc, "4.3 Monitoring và MLOps (1đ)", level=2)

    add_heading(doc, "4.3.1 Pipeline MLOps tự động", level=3)
    add_paragraph(doc, "Toàn bộ pipeline MLOps được tự động hóa:", bold=True)

    add_code_block(doc,
        "┌── Drift Detection (3 trigger sources) ──────────────────┐\n"
        "│ 1. Azure Function DriftMonitor (timer 1h)               │\n"
        "│ 2. GitHub Actions drift-detection.yml (cron 6h)         │\n"
        "│ 3. Manual: python ml/drift_monitor.py                   │\n"
        "└──────────────────────────────────────────────────────────┘\n"
        "                          ↓\n"
        "            ml/drift_monitor.py::run_monitor()\n"
        "              ├─ SQL: vw_ForecastVsActual\n"
        "              ├─ Compute: MAE, MAPE\n"
        "              ├─ Check: MAE > 25.0 AND samples ≥ 24 AND cooldown > 120min\n"
        "              └─ Lock: file-based, PID tracking, 600s timeout\n"
        "                          ↓ (if drift detected)\n"
        "        ┌─────────────────┼─────────────────┐\n"
        "        ↓                 ↓                 ↓\n"
        "   local mode         azureml mode     github_actions\n"
        "   retrain_and_       trigger_          dispatch\n"
        "   compare.py         training_         ci-cd-mlops.yml\n"
        "   --promote          pipeline.py\n"
        "        ↓                 ↓\n"
        "   Gate check:      AML Job → wait\n"
        "   new R² ≥ old?    → register model\n"
        "        ↓           → compare R² vs best\n"
        "   PROMOTE/REJECT   → promote/reject\n"
        "        ↓                 ↓\n"
        "        └────── Alerts ───┘\n"
        "               ├─ Slack webhook\n"
        "               ├─ Email SMTP\n"
        "               ├─ Teams Adaptive Card\n"
        "               └─ SQL: MonitoringEvents")

    add_heading(doc, "4.3.2 Model Registry (mlops/model_registry.py)", level=3)
    add_paragraph(doc,
        "Azure ML Model Registry qua azure-ai-ml SDK v2. "
        "Model name: 'sales-forecast-model'. Functions:")
    add_bullet(doc, "register_model(model_path, metrics) → đăng ký với tags (metric_mae, metric_r2, etc.)")
    add_bullet(doc, "list_versions(top_n) → liệt kê N version gần nhất")
    add_bullet(doc, "compare_versions(a, b) → so sánh metrics giữa 2 version")
    add_bullet(doc, "get_best_version(metric) → tìm version tốt nhất theo R² hoặc MAE")
    add_bullet(doc, "promote_model(version, stage) → gắn tag 'production', archive version cũ")

    add_heading(doc, "4.3.3 Health Check & Rollback (monitoring/model_health_check.py)", level=3)
    add_paragraph(doc,
        "Tự động backup model trước khi retrain. Nếu model mới hoạt động kém, "
        "rollback về version cũ từ rollback_backup/. "
        "Model files: revenue_model.pkl, quantity_model.pkl, label_encoders.pkl, model_metadata.json.")

    add_heading(doc, "4.3.4 A/B Shadow Testing (monitoring/ab_shadow_test.py)", level=3)
    add_paragraph(doc,
        "Route phần trăm traffic đến shadow model (mặc định 10%), "
        "log predictions vào shadow_predictions.jsonl. "
        "Sau thời gian đánh giá (mặc định 24h, tối thiểu 50 predictions), "
        "tự động promote hoặc reject shadow model.")

    add_heading(doc, "4.3.5 Telemetry (monitoring/telemetry.py)", level=3)
    add_paragraph(doc,
        "Application Insights SDK (opencensus). PipelineHealthMonitor theo dõi 7 thành phần: "
        "data_generator, event_hubs, azure_functions, stream_analytics, azure_sql, ml_endpoint, power_bi. "
        "Decorator @monitor_performance tự động đo thời gian và ghi dependency tracing.")

    add_heading(doc, "4.3.6 GitHub Actions CI/CD", level=3)
    add_paragraph(doc, "5 workflows:")
    add_table(doc,
        ["Workflow", "Trigger", "Chức năng"],
        [
            ["ci.yml", "Push/PR → main", "Lint, test, build"],
            ["ci-cd-mlops.yml", "Push main (terraform/ml/mlops paths) + workflow_call + manual",
             "Terraform apply → retrain → register → deploy endpoint"],
            ["drift-detection.yml", "Cron mỗi 6h + manual",
             "Preflight → drift-check → mae-check → auto-retrain (gọi ci-cd-mlops.yml)"],
            ["deploy-functions.yml", "—", "Deploy Azure Functions"],
            ["deploy-simulator.yml", "—", "Deploy data simulator"],
        ])

    add_heading(doc, "4.3.7 Security (security/key_vault.py)", level=3)
    add_paragraph(doc,
        "SecretManager class: ManagedIdentityCredential (Azure) / DefaultAzureCredential (local). "
        "Cache: @lru_cache(maxsize=32). Fallback: đọc từ environment variable nếu Key Vault không khả dụng.")
    add_paragraph(doc, "9 secrets được quản lý:", bold=True)
    add_bullet(doc, "event-hub-connection-string, sql-connection-string, sql-admin-password")
    add_bullet(doc, "ml-endpoint-url, ml-api-key, blob-connection-string")
    add_bullet(doc, "powerbi-push-url, appinsights-connection-string, openweather-api-key")

    doc.add_page_break()

    # ================================================================
    # CHƯƠNG 5: KẾT LUẬN
    # ================================================================
    add_heading(doc, "CHƯƠNG 5: KẾT LUẬN", level=1)

    add_heading(doc, "5.1 Kết quả đạt được", level=2)
    add_bullet(doc, "Xây dựng hệ thống phân tích bán hàng thời gian thực end-to-end trên Azure")
    add_bullet(doc, "15 tài nguyên Azure triển khai qua Terraform (IaC)")
    add_bullet(doc, "Pipeline real-time: Data Generator → Event Hub → Functions → Stream Analytics → SQL (latency trung bình 27.67ms)")
    add_bullet(doc, "Pipeline batch: ADF (4 activities, trigger hàng ngày 02:00 UTC)")
    add_bullet(doc, "ML pipeline: 9 mô hình so sánh, GradientBoosting (200 trees, depth 6) được chọn")
    add_bullet(doc, "MLOps tự động: drift detection (MAE-based) → retrain → register → deploy (blue/green)")
    add_bullet(doc, "Monitoring: DriftMonitor (1h), GitHub Actions (6h), health check, A/B shadow test")
    add_bullet(doc, "Dashboard: Power BI 5 trang + Flask Web App 10 routes + RLS 3 roles")
    add_bullet(doc, "Security: Key Vault (9 secrets), ManagedIdentity, TLS 1.2")
    add_bullet(doc, "Benchmark: 4.5 GB dataset, Parquet nhanh hơn CSV 31.87×")

    add_heading(doc, "5.2 Hạn chế", level=2)
    add_bullet(doc, "Cloud SQL chậm hơn local Parquet 147× cho batch analytics (do network latency)")
    add_bullet(doc, "Hiện deploy tại East US (290.73ms) — nên migrate sang Southeast Asia (69.28ms)")
    add_bullet(doc, "stream_query.sql thiếu HoppingWindow (chỉ có trong ARM template)")
    add_bullet(doc, "SVR training bị giới hạn 10,000 mẫu do tốc độ chậm")

    add_heading(doc, "5.3 Hướng phát triển", level=2)
    add_bullet(doc, "Migrate sang Azure region Southeast Asia để giảm latency")
    add_bullet(doc, "Đồng bộ stream_query.sql với ARM template (thêm HoppingWindow)")
    add_bullet(doc, "Thêm LightGBM/XGBoost vào compare_models.py (hiện đã có trong evaluate_pipelines.py)")
    add_bullet(doc, "Mở rộng số cửa hàng và sản phẩm cho production scale")

    doc.add_page_break()

    # ================================================================
    # TÀI LIỆU THAM KHẢO
    # ================================================================
    add_heading(doc, "TÀI LIỆU THAM KHẢO", level=1)
    refs = [
        "[1] Microsoft Azure Documentation – Event Hubs, Stream Analytics, SQL Database, ML",
        "[2] Scikit-learn Documentation – GradientBoostingRegressor, Model Selection",
        "[3] Terraform AzureRM Provider Documentation (~3.90)",
        "[4] Azure ML SDK v2 (azure-ai-ml) – Model Registry, Online Endpoints",
        "[5] Power BI DAX Reference – Time Intelligence, Row-Level Security",
        "[6] OpenCensus Python SDK – Application Insights Integration",
        "[7] GitHub Actions Documentation – Workflow syntax, Reusable workflows",
    ]
    for ref in refs:
        add_paragraph(doc, ref, size=11, space_after=3)

    # ── Save ──
    doc.save(REPORT_PATH)
    print(f"Report saved: {REPORT_PATH}")

    # Count
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    print(f"Paragraphs: {para_count}, Tables: {table_count}")


if __name__ == "__main__":
    build_report()
