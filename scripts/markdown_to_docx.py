#!/usr/bin/env python3
"""Minimal Markdown -> DOCX converter tailored for the project's thesis file.

Maps headings (#, ##, ###), bullets (- or *), blank lines (paragraph breaks),
and fenced code blocks (```) into a Word document using python-docx.

This intentionally keeps styling simple (Times New Roman, 12pt) and is
designed to run inside the existing venv where `python-docx` is already used
by `scripts/final_report.py`.
"""
import os
from docx import Document
from docx.shared import Pt

SRC = os.path.join(os.path.dirname(__file__), '..', 'docs', 'bao_cao_luan_van.md')
OUT = os.path.join(os.path.dirname(__file__), '..', 'docs', 'bao_cao_luan_van.docx')


def set_normal_font(run, size=12):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)


def convert(md_path, out_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    in_code = False
    code_lines = []
    para_lines = []

    def flush_paragraph():
        nonlocal para_lines
        if not para_lines:
            return
        text = ' '.join([l.strip() for l in para_lines]).strip()
        if text:
            p = doc.add_paragraph(text)
            set_normal_font(p.runs[0] if p.runs else p.add_run(), 12)
        para_lines = []

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        run = p.add_run('\n'.join(code_lines))
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        code_lines = []

    with open(md_path, encoding='utf-8') as f:
        for raw in f:
            line = raw.rstrip('\n')
            if line.strip().startswith('```'):
                if in_code:
                    # end code
                    flush_code()
                    in_code = False
                else:
                    # start code
                    flush_paragraph()
                    in_code = True
                continue

            if in_code:
                code_lines.append(line)
                continue

            # Headings
            if line.startswith('### '):
                flush_paragraph()
                level = 3
                h = doc.add_heading(line[4:].strip(), level=3)
                for run in h.runs:
                    run.font.name = 'Times New Roman'
                continue
            if line.startswith('## '):
                flush_paragraph()
                h = doc.add_heading(line[3:].strip(), level=2)
                for run in h.runs:
                    run.font.name = 'Times New Roman'
                continue
            if line.startswith('# '):
                flush_paragraph()
                h = doc.add_heading(line[2:].strip(), level=1)
                for run in h.runs:
                    run.font.name = 'Times New Roman'
                continue

            # Bullet
            stripped = line.lstrip()
            if stripped.startswith('- ') or stripped.startswith('* '):
                flush_paragraph()
                p = doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
                set_normal_font(p.runs[0] if p.runs else p.add_run(), 12)
                continue

            # Empty line => paragraph boundary
            if not line.strip():
                flush_paragraph()
                continue

            # Otherwise accumulate paragraph text
            para_lines.append(line)

    # flush any remaining
    flush_paragraph()
    if in_code:
        flush_code()

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    md = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'docs', 'bao_cao_luan_van.md'))
    out = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'docs', 'bao_cao_luan_van.docx'))
    if not os.path.exists(md):
        print('Source Markdown not found:', md)
        raise SystemExit(2)
    convert(md, out)
