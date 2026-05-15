"""Fix 4 discrepancies in Report.docx based on actual source code verification."""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

REPORT = "Report.docx"

def set_run_font(run, name="Times New Roman", size=13, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def fix_paragraph_text(para, old_text, new_text):
    """Replace text in paragraph while preserving formatting of first run."""
    full = para.text
    if old_text not in full:
        return False
    # Clear all runs, rewrite with same formatting as first run
    if para.runs:
        fmt = para.runs[0]
        font_name = fmt.font.name or "Times New Roman"
        font_size = fmt.font.size or Pt(13)
        is_bold = fmt.bold
    else:
        font_name, font_size, is_bold = "Times New Roman", Pt(13), False

    new_full = full.replace(old_text, new_text)
    # Remove existing runs
    for run in para.runs:
        run._element.getparent().remove(run._element)
    # Add new run
    run = para.add_run(new_full)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = is_bold
    return True


def main():
    doc = Document(REPORT)
    fixes = 0

    for i, p in enumerate(doc.paragraphs):
        t = p.text

        # ── FIX 1: PSI → MAE-based drift detection ──

        # Heading
        if t.strip() == "2.2.3. Drift Detection — Population Stability Index (PSI)":
            if fix_paragraph_text(p,
                "2.2.3. Drift Detection — Population Stability Index (PSI)",
                "2.2.3. Drift Detection — MAE-based Monitoring"):
                fixes += 1
                print(f"  Fixed para {i}: heading PSI → MAE")

        # Body para about PSI theory
        if "Hệ thống giám sát model drift bằng chỉ số PSI" in t:
            new_t = (
                "Hệ thống giám sát model drift bằng chỉ số MAE (Mean Absolute Error). "
                "Drift monitor so sánh MAE hiện tại của model với ngưỡng cho phép. "
                "Nếu MAE vượt ngưỡng (cho thấy model dự đoán kém hơn), hệ thống tự động "
                "trigger retraining pipeline qua Azure ML. "
                "Azure Function DriftMonitor chạy mỗi giờ, truy vấn SQL để tính MAE giữa "
                "dự báo và giá trị thực, ghi kết quả vào bảng log và gửi thông báo qua Slack."
            )
            # Clear and rewrite
            for run in p.runs:
                run._element.getparent().remove(run._element)
            run = p.add_run(new_t)
            set_run_font(run)
            fixes += 1
            print(f"  Fixed para {i}: PSI body → MAE body")

        # DriftMonitor description in Ch4
        if "Tính PSI giữa baseline" in t:
            if fix_paragraph_text(p,
                "Tính PSI giữa baseline và dữ liệu gần nhất, nếu PSI > 0.25 thì trigger retraining pipeline qua Azure ML",
                "Tính MAE giữa dự báo và giá trị thực, nếu MAE vượt ngưỡng thì trigger retraining pipeline qua Azure ML"):
                fixes += 1
                print(f"  Fixed para {i}: DriftMonitor PSI → MAE")

        # ── FIX 2: ADF pipeline — 6 activities → 4 activities ──

        # Service description
        if "CopyBlobToSQL → PrepareTrainingData → SubmitMLJob → WaitForMLJob → CheckMLSuccess → UpdateForecasts" in t:
            if fix_paragraph_text(p,
                "CopyBlobToSQL → PrepareTrainingData → SubmitMLJob → WaitForMLJob → CheckMLSuccess → UpdateForecasts",
                "CopyBlobToSQL → PrepareTrainingData → RunMLPipeline → UpdateForecasts"):
                fixes += 1
                print(f"  Fixed para {i}: ADF 6 activities → 4")

        # Detailed bullets in Ch3
        if "SubmitMLJob: WebActivity gọi Azure ML REST API" in t:
            if fix_paragraph_text(p,
                "SubmitMLJob: WebActivity gọi Azure ML REST API để submit training job.",
                "RunMLPipeline: AzureMLExecutePipeline activity gọi Azure ML pipeline để training model."):
                fixes += 1
                print(f"  Fixed para {i}: SubmitMLJob → RunMLPipeline")

        if "WaitForMLJob: Until loop polling" in t:
            if fix_paragraph_text(p,
                "WaitForMLJob: Until loop polling mỗi 60 giây chờ job hoàn thành.",
                "(Pipeline ADF chờ RunMLPipeline activity hoàn thành trước khi chuyển sang UpdateForecasts.)"):
                fixes += 1
                print(f"  Fixed para {i}: WaitForMLJob → note")

        if "CheckMLSuccess: IfCondition kiểm tra kết quả training" in t:
            if fix_paragraph_text(p,
                "CheckMLSuccess: IfCondition kiểm tra kết quả training (metrics threshold).",
                "(Kết quả training được kiểm tra bên trong Azure ML pipeline, model tự động register nếu đạt.)"):
                fixes += 1
                print(f"  Fixed para {i}: CheckMLSuccess → note")

        # ── FIX 3: SQL tables — remove Products/StoreRegions, fix SecurityMapping ──

        # Table in section 3.0 — the table rows are in doc.tables, handle separately

        # SecurityMapping reference in RLS section
        if "Bảng SecurityMapping trong SQL" in t:
            if fix_paragraph_text(p,
                "Bảng SecurityMapping trong SQL lưu mapping",
                "Bảng SecurityMapping (tạo qua script riêng) lưu mapping"):
                fixes += 1
                print(f"  Fixed para {i}: SecurityMapping clarification")

        if "sql/create_tables.sql (bảng SecurityMapping)" in t:
            if fix_paragraph_text(p,
                "sql/create_tables.sql (bảng SecurityMapping)",
                "powerbi/rls_config.dax"):
                fixes += 1
                print(f"  Fixed para {i}: SecurityMapping file ref")

        # ── FIX 4: 72.5M → ~70M rows ──
        if "72.500.000 rows" in t:
            if fix_paragraph_text(p, "72.500.000 rows", "~70.000.000 rows (target 4.5 GB)"):
                fixes += 1
                print(f"  Fixed para {i}: 72.5M → ~70M")

        if "72.5 triệu giao dịch" in t:
            if fix_paragraph_text(p, "72.5 triệu giao dịch", "~70 triệu giao dịch"):
                fixes += 1
                print(f"  Fixed para {i}: 72.5 → ~70 triệu")

    # ── FIX 3b: Fix table listing SQL tables ──
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells_text = [c.text.strip() for c in row.cells]
            # Find the row with Products / StoreRegions
            if "Products" in cells_text and "dimension" in " ".join(cells_text).lower():
                # Replace with WeatherSalesCorrelation
                for ci, cell in enumerate(row.cells):
                    if "Products" in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run._element.getparent().remove(run._element)
                            run = para.add_run("WeatherSalesCorrelation")
                            set_run_font(run, size=11)
                        fixes += 1
                        print(f"  Fixed table {ti} row {ri}: Products → WeatherSalesCorrelation")
                    if "dimension" in cell.text.lower():
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run._element.getparent().remove(run._element)
                            run = para.add_run("Tương quan thời tiết - doanh thu")
                            set_run_font(run, size=11)
                        fixes += 1
                    if "< 1 MB" in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run._element.getparent().remove(run._element)
                            run = para.add_run("~5 MB/tháng")
                            set_run_font(run, size=11)
                        fixes += 1

    doc.save(REPORT)
    print(f"\nDone: {fixes} fixes applied to {REPORT}")


if __name__ == "__main__":
    main()
