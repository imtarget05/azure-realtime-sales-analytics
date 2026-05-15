#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt
import os

md_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'docs', 'bao_cao_luan_van.md'))
out_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'docs', 'bao_cao_luan_van_updated.docx'))

def set_normal_font(run, size=12):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)

if not os.path.exists(md_path):
    print('Source not found', md_path)
    raise SystemExit(2)

with open(md_path, encoding='utf-8') as f:
    lines = f.readlines()

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

in_code = False
code_lines = []
para_lines = []

def flush_paragraph():
    global para_lines
    if not para_lines:
        return
    text = ' '.join([l.strip() for l in para_lines]).strip()
    if text:
        p = doc.add_paragraph(text)
        if p.runs:
            set_normal_font(p.runs[0], 12)
        else:
            set_normal_font(p.add_run(), 12)
    para_lines = []

def flush_code():
    global code_lines
    if not code_lines:
        return
    p = doc.add_paragraph()
    run = p.add_run('\n'.join(code_lines))
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    code_lines = []

for raw in lines:
    line = raw.rstrip('\n')
    if line.strip().startswith('```'):
        if in_code:
            flush_code()
            in_code = False
        else:
            flush_paragraph()
            in_code = True
        continue
    if in_code:
        code_lines.append(line)
        continue
    if line.startswith('### '):
        flush_paragraph()
        h = doc.add_heading(line[4:].strip(), level=3)
        for run in h.runs:
            run.font.name = 'Times New Roman'
        continue
    if line.startswith('## '):
        flush_paragraph()
        h = doc.add_heading(line[3:].strip(), level=2)
        for run in h.runs:
            run.font.name = 'Times New Roman'
        continue
    if line.startswith('# '):
        flush_paragraph()
        h = doc.add_heading(line[2:].strip(), level=1)
        for run in h.runs:
            run.font.name = 'Times New Roman'
        continue
    stripped = line.lstrip()
    if stripped.startswith('- ') or stripped.startswith('* '):
        flush_paragraph()
        p = doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
        if p.runs:
            set_normal_font(p.runs[0], 12)
        continue
    if not line.strip():
        flush_paragraph()
        continue
    para_lines.append(line)

flush_paragraph()
if in_code:
    flush_code()

try:
    doc.save(out_path)
    print('Saved:', out_path)
except Exception as e:
    print('Save failed:', e)
    raise
